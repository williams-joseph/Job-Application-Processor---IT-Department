"""
Document extraction module for ECOWAS Application Processor.
Handles PDF, DOCX, and OCR-based text extraction.
"""

import os
import re
import logging
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any
from datetime import datetime

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    import pytesseract
    from PIL import Image
except ImportError:
    pytesseract = None
    Image = None


from config import TESSERACT_PATH

if pytesseract and os.path.exists(TESSERACT_PATH):
    pytesseract.pytesseract.tesseract_cmd = TESSERACT_PATH


logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class FieldExtractor:
    """Extracts structured data from application forms."""
    
    # Regex patterns for field extraction
    PATTERNS = {
        'name': [
            r'(?:Full\s+)?Name\s*:?\s*([A-Z][a-zA-Z\s\'-]{2,}(?:\s+[A-Z][a-zA-Z\s\'-]{2,})+)',
            r'(?:First|Middle|Family)\s+Name\s*:?\s*([A-Z][a-zA-Z\s\'-]+)',
            r'Applicant\s+Name\s*:?\s*([A-Z][a-zA-Z\s\'-]+)',
            r'Surname\s*:?\s*([A-Z][a-zA-Z\s\'-]+)',
        ],
        'dob': [
            r'(?:Date\s+of\s+Birth|DOB|Birth\s+Date|Date\s+de\s+naissance|Data\s+nasc\.?)\s*:?\s*(\d{1,2}[-/\.\s]+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\w*[-/\.\s]+\d{2,4})',
            r'(?:Date\s+of\s+Birth|DOB|Birth\s+Date|Date\s+de\s+naissance|Data\s+nasc\.?)\s*:?\s*(\d{1,2}[-/\.]\d{1,2}[-/\.]\d{2,4})',
            r'(?:Date\s+of\s+Birth|DOB|Birth\s+Date|Date\s+de\s+naissance|Data\s+nasc\.?)\s*:?\s*([A-Z][a-z]+\s+\d{1,2},?\s+\d{4})',
        ],
        'nationality': [
            r'(?:Nationality|Nationalité|Nacionalidade)\s*:?\s*([A-Z][a-zA-Z\s]+?)(?:\n|Sex|Gender|Date|Sexe|Sexo|Etat|Estado)',
            r'Country\s+of\s+Citizenship\s*:?\s*([A-Z][a-zA-Z\s]+)',
        ],
        'gender': [
            r'(?:Sex|Gender|Sexe|Sexo)\s*:?\s*(Male|Female|M|F|Masculin|Féminin|Masculino|Feminino|Homme|Mulher)\b',
        ],
        'exp_start': [
            r'(?:Experience\s+Start|Work\s+Start|Started\s+Work)\s*:?\s*(\d{4})',
            r'(?:Earliest|First)\s+Employment\s*:?\s*(\d{4})',
            r'(?:Experience|Employment)\s+History\s+since\s*:?\s*(\d{4})',
        ]
    }
    
    def __init__(self):
        self.confidence_threshold = 0.5
        from config import DEFAULT_POSITION_CODE, DEFAULT_INT_EXT
        self.default_position_code = DEFAULT_POSITION_CODE
        self.default_int_ext = DEFAULT_INT_EXT
    
    def extract_from_file(self, file_path: str) -> Dict:
        """
        Extract text and parse fields from a document file.
        
        Args:
            file_path: Path to the application form
            
        Returns:
            Dictionary with extracted fields and metadata
        """
        file_path = Path(file_path)
        
        result = {
            'file_path': str(file_path),
            'file_name': file_path.name,
            'extraction_status': 'success',
            'error_message': None,
            'fields': {},
            'confidence_scores': {},
            'raw_text': None,
        }
        
        try:
            # Extract text based on file type
            structured = None
            if file_path.suffix.lower() == '.pdf':
                text = self._extract_pdf(file_path)
            elif file_path.suffix.lower() in ['.docx', '.doc']:
                docx_data = self._extract_docx(file_path)
                text = docx_data['text']
                structured = docx_data.get('structured')
            elif file_path.suffix.lower() in ['.jpg', '.jpeg', '.png', '.tiff', '.bmp']:
                text = self._extract_image(file_path)
            else:
                result['extraction_status'] = 'unsupported'
                result['error_message'] = f"Unsupported file type: {file_path.suffix}"
                return result
            
            if not text or len(text.strip()) < 10:
                result['extraction_status'] = 'failed'
                result['error_message'] = "No text extracted from document"
                return result
            
            result['raw_text'] = text
            
            # Parse fields from extracted text
            parsed_fields = self._parse_fields(text, structured)
            result['fields'] = parsed_fields['fields']
            result['confidence_scores'] = parsed_fields['confidence']
            result['errors'] = parsed_fields.get('errors', [])
            
        except Exception as e:
            result['extraction_status'] = 'error'
            result['error_message'] = str(e)
        
        return result
    
    def _extract_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file."""
        if not pdfplumber:
            raise ImportError("pdfplumber not installed")
        
        text = ""
        from config import PDF_MAX_PAGES
        
        try:
            with pdfplumber.open(file_path) as pdf:
                num_pages = len(pdf.pages)
                pages_to_process = min(num_pages, PDF_MAX_PAGES)
                
                logger.info(f"Extracting text from PDF: {file_path.name} ({pages_to_process}/{num_pages} pages)")
                
                for i in range(pages_to_process):
                    page = pdf.pages[i]
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            logger.error(f"Error reading PDF {file_path.name}: {e}")
            return ""
        
        # If no text extracted, try OCR
        if not text.strip() and pytesseract:
            logger.info(f"No text found in {file_path.name}, attempting OCR...")
            text = self._extract_pdf_with_ocr(file_path)
        
        return text
    
    def _extract_pdf_with_ocr(self, file_path: Path) -> str:
        """Extract text from PDF using OCR."""
        if not pytesseract or not Image:
            return ""
        
        text = ""
        from config import PDF_MAX_PAGES
        
        try:
            with pdfplumber.open(file_path) as pdf:
                num_pages = len(pdf.pages)
                pages_to_process = min(num_pages, PDF_MAX_PAGES)
                
                for i in range(pages_to_process):
                    logger.info(f"OCR-ing page {i+1}/{pages_to_process} of {file_path.name}...")
                    page = pdf.pages[i]
                    # Convert page to image
                    img = page.to_image(resolution=300)
                    pil_img = img.original
                    # OCR the image with timeout (requires Tesseract 4.0+)
                    try:
                        page_text = pytesseract.image_to_string(pil_img, timeout=60)
                        text += page_text + "\n"
                    except RuntimeError as re:
                         logger.warning(f"OCR timeout/error on page {i+1} of {file_path.name}: {re}")
        except Exception as e:
            logger.error(f"OCR failed for {file_path.name}: {e}")
            
        return text
    
    def _extract_docx(self, file_path: Path) -> Dict:
        """Extract text and structured data from DOCX file."""
        if not Document:
            raise ImportError("python-docx not installed")
        
        doc = Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        
        # Also extract text from tables for legacy pattern matching
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += " " + cell.text
            text += "\n"
            
        # Specific table extraction logic
        structured_data = {
            'personal_info': {},
            'education': [],
            'qualifications': [],
            'experience': [] # Added initialization for experience
        }
        
        for table in doc.tables:
            mode = "personal" # Start with personal info detection
            for row in table.rows:
                try:
                    unique_cells = self._get_unique_cells(row)
                    if not unique_cells: continue
                    row_text_all = " ".join(unique_cells).lower()
                    
                    # Mode switching with stronger signals
                    if any(k in row_text_all for k in ["level of education", "niveau d’éducation", "nível de educação"]):
                        mode = "education"
                        continue 
                    elif any(k in row_text_all for k in ["title of qualification", "titre de la qualification", "título de qualificação"]):
                        mode = "qual"
                        continue
                    elif any(k in row_text_all for k in ["starting date", "date de début", "data de início"]):
                        mode = "exp"
                        continue
                    elif any(k in row_text_all for k in ["professional experience", "expériences professionnelles", "experiência profissional"]):
                        mode = "exp"
                        continue
                    elif any(k in row_text_all for k in ["professional references", "références professionnelles"]):
                        mode = None
                        continue
                    elif any(k in row_text_all for k in ["personal information", "informations générales", "informações pessoais"]):
                        mode = "personal"
                        continue

                    if mode == "personal":
                        # Handle empty first cell if it exists
                        label_idx = 0 if unique_cells[0] != "" else (1 if len(unique_cells) > 1 else 0)
                        val_idx = label_idx + 1
                        
                        label = unique_cells[label_idx].lower()
                        val = unique_cells[val_idx] if len(unique_cells) > val_idx else ""
                        
                        if any(k in label for k in ["first name", "prénoms"]): structured_data['personal_info']['first'] = val
                        elif any(k in label for k in ["second name", "autre nom", "autrenoms"]): structured_data['personal_info']['second'] = val
                        elif any(k in label for k in ["family name", "nom de famille"]): structured_data['personal_info']['family'] = val
                        elif any(k in label for k in ["gender", "sexe", "sexo"]): structured_data['personal_info']['gender'] = val
                        elif any(k in label for k in ["nationality", "nationalité", "nacionalidade"]): structured_data['personal_info']['nationality'] = val
                        elif any(k in label for k in ["date of birth", "naissance", "data nasc"]): structured_data['personal_info']['dob'] = val
                        
                    elif mode == "education":
                        idx_offset = 0
                        if re.match(r'^\d+\.?\s*$', unique_cells[0]): idx_offset = 1
                        elif unique_cells[0] == '' and len(unique_cells) > 1 and re.match(r'^\d+\.?\s*$', unique_cells[1]): idx_offset = 2
                        elif unique_cells[0] == '': idx_offset = 1
                        
                        if len(unique_cells) >= 3 + idx_offset:
                            level = unique_cells[idx_offset]
                            year = unique_cells[idx_offset+1]
                            field = unique_cells[idx_offset+2] if len(unique_cells) > idx_offset+2 else ""
                            inst = unique_cells[idx_offset+3] if len(unique_cells) > idx_offset+3 else ""
                            
                            if level and level.strip() and not any(k in level.lower() for k in ["education", "niveau", "diplôme", "habilitação"]):
                                structured_data['education'].append({
                                    'level': level.strip(), 'year': year.strip(), 'field': field.strip(), 'inst': inst.strip()
                                })
                                
                    elif mode == "qual":
                        idx_offset = 1 if unique_cells[0] == '' else 0
                        if len(unique_cells) >= 1 + idx_offset:
                            title = unique_cells[idx_offset]
                            centre = unique_cells[idx_offset+1] if len(unique_cells) > idx_offset+1 else ""
                            year = unique_cells[idx_offset+2] if len(unique_cells) > idx_offset+2 else ""
                            
                            if title and title.strip() and not any(k in title.lower() for k in ["qualification", "titre", "certificat", "year", "année"]):
                                structured_data['qualifications'].append({
                                    'title': title.strip(), 'centre': centre.strip(), 'year': year.strip()
                                })

                    elif mode == "exp":
                        # Detect if it's a data row and not a header
                        idx_offset = 1 if unique_cells[0] == '' else 0
                        if len(unique_cells) >= 2 + idx_offset:
                            start = unique_cells[idx_offset]
                            end = unique_cells[idx_offset+1]
                            
                            # Clean up start (remove leading S/N)
                            start = re.sub(r'^\d+\.?\s*', '', start).strip()
                            
                            if start and not any(k in start.lower() for k in ["starting", "début", "início", "experience", "date", "fin", "month"]):
                                structured_data['experience'].append({'start': start, 'end': end})
                except:
                    continue
                
        return {
            'text': text,
            'structured': structured_data
        }

    def _get_unique_cells(self, row):
        """Get unique cell values from a row (de-duplicating merged cells)."""
        cells = [c.text.strip() for c in row.cells]
        unique_cells = []
        if cells:
            unique_cells.append(cells[0])
            for j in range(1, len(cells)):
                if cells[j] != cells[j-1]:
                    unique_cells.append(cells[j])
        return unique_cells
    
    def _extract_image(self, file_path: Path) -> str:
        """Extract text from image using OCR."""
        if not pytesseract or not Image:
            raise ImportError("pytesseract and Pillow not installed")
        
        img = Image.open(file_path)
        text = pytesseract.image_to_string(img)
        return text
    
    def _parse_fields(self, text: str, structured: Optional[Dict] = None) -> Dict:
        """Parse structured fields from text or DOCX data."""
        today = datetime.now()
        current_year = today.year
        
        fields = {
            'NAME': '',
            'POSITION CODE': self.default_position_code,
            'GENDER': '',
            'INT/EXT': self.default_int_ext,
            'DOB': '',
            'AGE': '',
            'NATIONALITY': '',
            'EXP START (YEAR)': '',
            'EXPERIENCE(Years)': '',
            'QUALIFICATIONS': '',
        }
        
        confidence = {k: 0.0 for k in fields.keys()}
        confidence['POSITION CODE'] = 1.0
        confidence['INT/EXT'] = 1.0
        
        extraction_errors = []

        # 1. Use structured data first (DOCX)
        if structured:
            p_info = structured.get('personal_info', {})
            
            # Name
            first = p_info.get('first', '').strip()
            second = p_info.get('second', '').strip()
            family = p_info.get('family', '').strip()
            if first or family:
                fields['NAME'] = f"{first} {second} {family}".replace('  ', ' ').strip().upper()
                confidence['NAME'] = 1.0
            
            # Gender
            gender_val = p_info.get('gender', '').strip().upper()
            if any(k in gender_val for k in ['MALE', 'MASCULIN', 'HOMME']): fields['GENDER'] = 'M'
            elif any(k in gender_val for k in ['FEMALE', 'FÉMININ', 'MULHER', 'FEMININO']): fields['GENDER'] = 'F'
            else: fields['GENDER'] = gender_val[:1].upper() if gender_val else ''
            if fields['GENDER']: confidence['GENDER'] = 1.0
            
            # DOB
            fields['DOB'] = p_info.get('dob', '').strip()
            if fields['DOB']: confidence['DOB'] = 1.0
            
            # Nationality
            fields['NATIONALITY'] = p_info.get('nationality', '').strip()
            if fields['NATIONALITY']: confidence['NATIONALITY'] = 1.0
            
            # Qualifications - Collect All
            raw_quals = []
            for edu in structured.get('education', []):
                raw_quals.append(f"{edu['level']}, {edu['field']}, {edu['inst']} - {edu['year']}")
            for cert in structured.get('qualifications', []):
                raw_quals.append(f"{cert['title']}, {cert['centre']} - {cert['year']}")
            
            # Deduplicate
            all_quals_set = set()
            for q in raw_quals:
                clean_q = q.strip()
                if clean_q and len(clean_q) > 10: all_quals_set.add(clean_q)
            
            all_quals = list(all_quals_set)
            
            def extract_sort_year(q_str):
                match = re.search(r'(\d{4})', q_str.split(' - ')[-1]) if ' - ' in q_str else None
                if not match: match = re.search(r'(\d{4})', q_str)
                return int(match.group(1)) if match else 0

            all_quals.sort(key=extract_sort_year, reverse=True)
            fields['QUALIFICATIONS'] = "\n".join(all_quals)
            if fields['QUALIFICATIONS']: confidence['QUALIFICATIONS'] = 1.0

            # Experience calculation
            exp_entries = structured.get('experience', [])
            total_months = 0
            start_years = []

            for entry in exp_entries:
                try:
                    start_dt = self._parse_month_year(entry['start'])
                    end_text = entry['end'].strip().lower() if entry['end'] else ""
                    if any(k in end_text for k in ['present', 'date', 'current', 'até agora']):
                        end_dt = today
                    else:
                        end_dt = self._parse_month_year(entry['end'])
                    
                    if start_dt:
                        start_years.append(start_dt.year)
                        if end_dt:
                            months = (end_dt.year - start_dt.year) * 12 + (end_dt.month - start_dt.month)
                            if months > 0: total_months += months
                except:
                    continue

            if start_years:
                fields['EXP START (YEAR)'] = min(start_years)
                confidence['EXP START (YEAR)'] = 1.0
            
            if total_months > 0:
                fields['EXPERIENCE(Years)'] = round(total_months / 12, 1)
                confidence['EXPERIENCE(Years)'] = 1.0

        # 2. Fallbacks and normalization
        
        # Name fallback
        if not fields['NAME']:
            name_match = self._find_best_match(text, self.PATTERNS['name'])
            if name_match:
                match_val = name_match[0].strip().upper()
                if len(match_val) > 2 and match_val not in ["PR", "S/N"]:
                    fields['NAME'] = match_val
                    confidence['NAME'] = 0.9

        # DOB Fallback & Standardize
        if not fields['DOB']:
            dob_match = self._find_best_match(text, self.PATTERNS['dob'])
            if dob_match: 
                fields['DOB'] = dob_match[0].strip()
                confidence['DOB'] = 0.5
        
        if fields['DOB']:
            normalized_dob = self._normalize_date(fields['DOB'])
            if normalized_dob:
                try:
                    birth_date = datetime.strptime(normalized_dob, '%Y-%m-%d')
                    fields['AGE'] = today.year - birth_date.year - ((today.month, today.day) < (birth_date.month, birth_date.day))
                    # Standardize to "01 January 1970"
                    fields['DOB'] = birth_date.strftime('%d %B %Y')
                    confidence['AGE'] = max(confidence.get('AGE', 0), 0.9)
                except:
                    pass
            else:
                # If cannot normalize, leave as is but log that it might be invalid
                extraction_errors.append(f"Cell for DOB is left original because tool couldn't standardize format: '{fields['DOB']}'")

        # Nationality & Gender fallback
        if not fields['NATIONALITY']:
            nat_match = self._find_best_match(text, self.PATTERNS['nationality'])
            if nat_match: fields['NATIONALITY'] = nat_match[0].strip()
            
        if not fields['GENDER']:
            gender_match = self._find_best_match(text, self.PATTERNS['gender'])
            if gender_match:
                gv = gender_match[0].strip().upper()
                if any(k in gv for k in ['MALE', 'MASCULIN', 'HOMME']): fields['GENDER'] = 'M'
                elif any(k in gv for k in ['FEMALE', 'FÉMININ', 'MULHER']): fields['GENDER'] = 'F'
                else: fields['GENDER'] = gv[:1]

        # Final cleanup: Replace any noise or empty placeholders with empty strings
        noise_keywords = [
            "PR", "S/N", "N/A", "NONE", "NO", "NULL", "UNDEFINED", "EMPTY", "SPECIFY", 
            "CLICK HERE", "AUTRENOMS", "FIRST NAME", "NOM DE FAMILLE", "FAMILY NAME",
            "SURNAME", "NAME:", "PRÉNOMS", "NOM", "AUTRE NOM"
        ]
        for k in fields:
            val_str = str(fields[k]).strip().upper()
            if val_str in noise_keywords or len(val_str) == 0:
                fields[k] = ""
            # Prevent partial noise like "S/N 1"
            if val_str.startswith("S/N "):
                fields[k] = ""
            
        # Error logging for any missing fields
        field_labels = {
            'NAME': 'Applicant Name',
            'GENDER': 'Gender',
            'DOB': 'Date of Birth',
            'AGE': 'Age',
            'NATIONALITY': 'Nationality',
            'EXP START (YEAR)': 'Experience Start Year',
            'EXPERIENCE(Years)': 'Total Years of Experience',
            'QUALIFICATIONS': 'Qualifications/Education'
        }
        
        for key, label in field_labels.items():
            val = fields.get(key)
            if val is None or str(val).strip() == "":
                extraction_errors.append(f"{label} is empty or could not be extracted. Please review this applicant manually.")

        return {
            'fields': fields,
            'confidence': confidence,
            'errors': extraction_errors
        }

    def _parse_month_year(self, date_str: str) -> Optional[datetime]:
        """Parse 'Month Year' formats."""
        if not date_str: return None
        # Clean string
        ds = re.sub(r'[^\w\s]', ' ', date_str).strip()
        fmts = ['%B %Y', '%b %Y', '%m %Y', '%Y']
        for f in fmts:
            try: return datetime.strptime(ds, f)
            except: continue
        # Try finding year
        match = re.search(r'(\d{4})', date_str)
        if match: return datetime(int(match.group(1)), 1, 1)
        return None

    def _extract_multiline_qualifications(self, text: str) -> str:
        """
        Extract qualifications and format them as:
        Masters Degree MBA - 2023
        B.sc Business Administration - 2017
        ...
        """
        qual_list = []
        
        # Keywords for degrees and diplomas
        degree_keywords = [
            'PhD', 'Doctorate', 'Master', 'Masters', 'MBA', 'MSc', 'MA', 
            'Bachelor', 'B.sc', 'BSc', 'B.A', 'BA', 'Degree',
            'Diploma', 'HND', 'OND', 'SSCE', 'WAEC', 'WEAC', 
            'School Certificate', 'Certificate'
        ]
        
        # Pattern to find a degree and a year near it
        # Example: "B.sc Business Administration (2017)" or "MSc MBA - 2023"
        # We'll look for lines or segments containing keywords and a 4-digit year
        
        lines = text.split('\n')
        for line in lines:
            line = line.strip()
            found_year = re.search(r'\b(19|20)\d{2}\b', line)
            if found_year:
                year = found_year.group(0)
                # Check if this line contains a degree keyword
                for keyword in degree_keywords:
                    if re.search(rf'\b{re.escape(keyword)}\b', line, re.IGNORECASE):
                        # Clean up the line to just the qualification part
                        clean_line = line.replace(year, '').strip()
                        # Remove common separators and brackets
                        clean_line = re.sub(r'[\(\)\-\:\,]', ' ', clean_line).strip()
                        # Re-format properly
                        qual_list.append(f"{clean_line} - {year}")
                        break
        
        # Remove duplicates and noise
        unique_quals_set = set()
        for q in qual_list:
            clean_q = q.strip()
            # Filter out short noise or duplicates
            if clean_q and len(clean_q) > 10:
                unique_quals_set.add(clean_q)
        
        unique_quals = list(unique_quals_set)
        
        # Sort by year descending
        try:
            def get_year(x):
                m = re.search(r'(\d{4})$', x)
                return int(m.group(1)) if m else 0
            unique_quals.sort(key=get_year, reverse=True)
        except Exception:
            pass
            
        return '\n'.join(unique_quals)
    
    def _find_best_match(self, text: str, patterns: list) -> Optional[Tuple]:
        """Try multiple patterns and return the best match."""
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
            if match:
                return match.groups()
        return None
    
    def _normalize_date(self, date_str: str) -> Optional[str]:
        """Normalize date to standard format (YYYY-MM-DD), handling multilingual inputs."""
        if not date_str: return None
        
        ds = date_str.lower().strip()
        
        # 0. Clean common punctuations and ordinals
        ds = re.sub(r'(\d+)(st|nd|rd|th)', r'\1', ds) # 20th -> 20
        ds = ds.replace(',', ' ').replace('.', ' ').replace('/', ' ').replace('-', ' ').replace(' de ', ' ')
        ds = re.sub(r'\s+', ' ', ds).strip()

        # 1. Map months (English, French, Portuguese)
        months = {
            'jan': 1, 'fev': 2, 'mar': 3, 'apr': 4, 'avr': 4, 'abr': 4, 'may': 5, 'mai': 5, 'jun': 6,
            'jul': 7, 'aug': 8, 'aou': 8, 'ago': 8, 'sep': 9, 'set': 9, 'oct': 10, 'out': 10, 
            'nov': 11, 'dec': 12, 'dez': 12, 
            'january':1, 'february':2, 'march':3, 'april':4, 'may':5, 'june':6, 'july':7, 'august':8, 
            'september':9, 'october':10, 'november':11, 'december':12,
            'janvier':1, 'fevrier':2, 'mars':3, 'avril':4, 'mai':5, 'juin':6, 'juillet':7, 'aout':8, 
            'septembre':9, 'octobre':10, 'novembre':11, 'decembre':12,
            'janeiro':1, 'fevereiro':2, 'marco':3, 'abril':4, 'maio':5, 'junho':6, 'julho':7, 'agosto':8, 
            'setembro':9, 'outubro':10, 'novembro':11, 'dezembro':12
        }
        
        # Try to find text month
        for name, num in months.items():
            if name in ds:
                # Replace name with number
                ds = ds.replace(name, str(num))
                break
        
        # Standard formats
        date_formats = [
            '%d %m %Y', '%Y %m %d', '%m %d %Y', 
            '%d %B %Y', '%d %b %Y' # fallbacks
        ]
        
        for fmt in date_formats:
            try:
                dt = datetime.strptime(ds, fmt)
                return dt.strftime('%Y-%m-%d')
            except ValueError:
                continue
        
        # Last resort: Try parsing component by component
        parts = ds.split()
        if len(parts) == 3:
            try:
                # Assumption: If year is first (YYYY MM DD), else (DD MM YYYY)
                p1, p2, p3 = int(parts[0]), int(parts[1]), int(parts[2])
                if p1 > 1900: # YYYY MM DD
                   return datetime(p1, p2, p3).strftime('%Y-%m-%d')
                else: # DD MM YYYY
                   return datetime(p3, p2, p1).strftime('%Y-%m-%d')
            except:
                pass

        return None
    
    def validate_fields(self, fields: Dict) -> Dict[str, bool]:
        """
        Validate extracted fields.
        
        Returns:
            Dictionary mapping field names to validation status
        """
        validation = {}
        
        # Name validation
        validation['NAME'] = bool(fields.get('NAME')) and len(str(fields['NAME']).split()) >= 2
        
        # DOB validation
        dob = fields.get('DOB', '')
        validation['DOB'] = bool(dob) and (
            re.match(r'\d{4}-\d{2}-\d{2}', str(dob)) or
            re.match(r'\d{1,2}[-/\.]\d{1,2}[-/\.]\d{2,4}', str(dob))
        )
        
        # Qualification validation
        validation['QUALIFICATIONS'] = bool(fields.get('QUALIFICATIONS'))
        
        # Nationality validation
        validation['NATIONALITY'] = bool(fields.get('NATIONALITY'))
        
        # Gender validation
        gender = fields.get('GENDER', '')
        validation['GENDER'] = str(gender) in ['Male', 'Female', 'M', 'F']
        
        return validation
