#!/usr/bin/env python3
"""
Test script to validate the extractor functionality.
Creates sample test files and runs extraction.
"""

import os
from pathlib import Path
from docx import Document
from docx.shared import Pt
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


def create_test_structure():
    """Create test folder structure with sample application forms."""
    
    base_dir = Path("test_data")
    base_dir.mkdir(exist_ok=True)
    
    # Sample applicants
    applicants = [
        {
            'name': 'John Doe',
            'dob': '15/03/1990',
            'qualification': 'BSc in Computer Science',
            'nationality': 'Nigerian',
            'sex': 'Male',
        },
        {
            'name': 'Jane Smith',
            'dob': '22/07/1988',
            'qualification': 'MSc in Information Technology',
            'nationality': 'Ghanaian',
            'sex': 'Female',
        },
        {
            'name': 'Abdul Rahman',
            'dob': '10/11/1992',
            'qualification': 'BSc in Software Engineering',
            'nationality': 'Nigerian',
            'sex': 'Male',
        },
        {
            'name': 'Mary Johnson',
            'dob': '05/01/1995',
            'qualification': 'HND in Computer Science',
            'nationality': 'Togolese',
            'sex': 'Female',
        },
        {
            'name': 'Emmanuel Okafor',
            'dob': '18/09/1987',
            'qualification': 'MSc in Cybersecurity',
            'nationality': 'Nigerian',
            'sex': 'Male',
        },
    ]
    
    logger.info(f"Creating test structure in {base_dir}")
    
    for idx, applicant in enumerate(applicants):
        # Create applicant folder
        applicant_folder = base_dir / applicant['name']
        applicant_folder.mkdir(exist_ok=True)
        
        # Create application form (DOCX)
        form_path = applicant_folder / "Application Form.docx"
        create_application_form(form_path, applicant)
        
        # Create additional dummy files
        (applicant_folder / "Resume.pdf").touch()
        (applicant_folder / "Certificate.pdf").touch()
        
        logger.info(f"Created folder for {applicant['name']}")
    
    logger.info(f"Test structure created successfully!")
    logger.info(f"Total applicants: {len(applicants)}")
    logger.info(f"Location: {base_dir.absolute()}")
    
    return base_dir


def create_application_form(file_path: Path, data: dict):
    """Create a sample application form in DOCX format."""
    
    doc = Document()
    
    # Title
    title = doc.add_heading('ECOWAS JOB APPLICATION FORM', 0)
    title.alignment = 1  # Center
    
    doc.add_paragraph()
    
    # Personal Information Section
    doc.add_heading('PERSONAL INFORMATION', level=1)
    
    # Add fields
    p = doc.add_paragraph()
    p.add_run('Full Name: ').bold = True
    p.add_run(data['name'])
    
    p = doc.add_paragraph()
    p.add_run('Date of Birth: ').bold = True
    p.add_run(data['dob'])
    
    p = doc.add_paragraph()
    p.add_run('Sex: ').bold = True
    p.add_run(data['sex'])
    
    p = doc.add_paragraph()
    p.add_run('Nationality: ').bold = True
    p.add_run(data['nationality'])
    
    doc.add_paragraph()
    
    # Educational Background Section
    doc.add_heading('EDUCATIONAL BACKGROUND', level=1)
    
    p = doc.add_paragraph()
    p.add_run('Highest Qualification: ').bold = True
    p.add_run(data['qualification'])
    
    doc.add_paragraph()
    
    # Additional sections (to make it realistic)
    doc.add_heading('WORK EXPERIENCE', level=1)
    doc.add_paragraph('Previous employer: Tech Solutions Ltd.')
    doc.add_paragraph('Duration: 2015 - 2020')
    doc.add_paragraph('Position: Software Developer')
    
    doc.add_paragraph()
    
    doc.add_heading('DECLARATION', level=1)
    doc.add_paragraph(
        'I hereby declare that the information provided above is true and correct '
        'to the best of my knowledge.'
    )
    
    doc.add_paragraph()
    p = doc.add_paragraph('Signature: ___________________')
    p = doc.add_paragraph(f'Date: {data["dob"]}')
    
    # Save document
    doc.save(file_path)


def test_extraction():
    """Test the extraction functionality."""
    try:
        from scanner import FolderScanner
        from extractor import FieldExtractor
        
        # Create test structure
        test_dir = create_test_structure()
        
        logger.info("\n" + "="*60)
        logger.info("Testing Extraction Functionality")
        logger.info("="*60 + "\n")
        
        # Initialize scanner and extractor
        scanner = FolderScanner()
        extractor = FieldExtractor()
        
        # Scan folders
        logger.info("Scanning folders...")
        applicants = scanner.scan_folders(str(test_dir))
        logger.info(f"Found {len(applicants)} applicant folders\n")
        
        # Test extraction on first applicant
        if applicants:
            test_applicant = applicants[0]
            logger.info(f"Testing extraction on: {test_applicant['applicant_name']}")
            logger.info(f"Form path: {test_applicant['application_form']}\n")
            
            if test_applicant['application_form']:
                result = extractor.extract_from_file(test_applicant['application_form'])
                
                logger.info("Extraction Results:")
                logger.info(f"Status: {result['extraction_status']}")
                logger.info(f"\nExtracted Fields:")
                for field, value in result['fields'].items():
                    logger.info(f"  {field}: {value}")
                
                logger.info(f"\nValidation:")
                validation = extractor.validate_fields(result['fields'])
                for field, valid in validation.items():
                    status = "✓" if valid else "✗"
                    logger.info(f"  {status} {field}")
        
        logger.info("\n" + "="*60)
        logger.info("Test Complete!")
        logger.info("="*60)
        logger.info(f"\nTest data location: {test_dir.absolute()}")
        logger.info("You can now run the main application and select this folder.")
        
    except ImportError as e:
        logger.error(f"Import error: {e}")
        logger.error("Make sure all dependencies are installed:")
        logger.error("  pip install -r requirements.txt")
    except Exception as e:
        logger.error(f"Test failed: {e}")
        import traceback
        traceback.print_exc()


if __name__ == '__main__':
    test_extraction()
